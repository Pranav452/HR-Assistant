"""
Document processing service for extracting and chunking text
"""

import os
import asyncio
from typing import List, Dict, Any
import pdfplumber
from docx import Document
import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
import hashlib
from datetime import datetime

from app.core.config import settings
from app.models.schemas import DocumentChunk

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    async def process_document(self, file_path: str, filename: str) -> List[DocumentChunk]:
        """Process a document and return chunks"""
        try:
            # Extract text based on file type
            text = await self._extract_text(file_path)
            
            # Create metadata
            metadata = {
                "filename": filename,
                "file_path": file_path,
                "processed_date": datetime.now().isoformat(),
                "file_size": os.path.getsize(file_path),
                "category": self._categorize_document(filename, text)
            }
            
            # Split text into chunks
            chunks = await self._create_chunks(text, metadata)
            
            return chunks
            
        except Exception as e:
            raise Exception(f"Error processing document {filename}: {str(e)}")
    
    async def _extract_text(self, file_path: str) -> str:
        """Extract text from different file formats"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return await self._extract_pdf_text(file_path)
        elif file_extension in ['.docx', '.doc']:
            return await self._extract_docx_text(file_path)
        elif file_extension == '.txt':
            return await self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n[Page {page_num + 1}]\n{page_text}\n"
        except Exception as e:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n[Page {page_num + 1}]\n{page_text}\n"
        
        return text.strip()
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        return text.strip()
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    async def _create_chunks(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into chunks with metadata"""
        # Create LangChain document
        doc = LangChainDocument(page_content=text, metadata=metadata)
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        # Convert to DocumentChunk objects
        document_chunks = []
        for i, chunk in enumerate(chunks):
            chunk_id = hashlib.md5(
                f"{metadata['filename']}_{i}_{chunk.page_content[:100]}".encode()
            ).hexdigest()
            
            chunk_metadata = {
                **metadata,
                "chunk_index": i,
                "chunk_id": chunk_id,
                "chunk_size": len(chunk.page_content)
            }
            
            document_chunks.append(DocumentChunk(
                content=chunk.page_content,
                metadata=chunk_metadata,
                chunk_id=chunk_id
            ))
        
        return document_chunks
    
    def _categorize_document(self, filename: str, text: str) -> str:
        """Categorize document based on filename and content"""
        filename_lower = filename.lower()
        text_lower = text.lower()
        
        # Category keywords
        categories = {
            'benefits': ['benefit', 'insurance', 'health', 'dental', 'vision', '401k', 'retirement'],
            'leave-policies': ['leave', 'vacation', 'pto', 'sick', 'maternity', 'paternity', 'fmla'],
            'work-policies': ['remote', 'work from home', 'wfh', 'flexible', 'schedule', 'attendance'],
            'performance': ['performance', 'review', 'evaluation', 'appraisal', 'feedback', 'goals'],
            'conduct': ['conduct', 'behavior', 'harassment', 'discrimination', 'ethics', 'compliance'],
            'compensation': ['salary', 'wage', 'pay', 'compensation', 'bonus', 'raise', 'promotion']
        }
        
        # Check filename first
        for category, keywords in categories.items():
            if any(keyword in filename_lower for keyword in keywords):
                return category
        
        # Check content
        for category, keywords in categories.items():
            keyword_count = sum(1 for keyword in keywords if keyword in text_lower)
            if keyword_count >= 2:  # At least 2 keywords found
                return category
        
        return 'general'
